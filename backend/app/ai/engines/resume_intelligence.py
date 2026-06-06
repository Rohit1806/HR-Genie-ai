"""
AI ENGINE 1: Resume Intelligence
File: backend/app/ai/engines/resume_intelligence.py

Extracts structured data from resumes (PDF/DOCX):
- Personal info, skills, education, experience
- Computes years of experience
- Normalizes skill names
- Returns clean JSON ready for DB storage
"""

import re
import logging
from pathlib import Path
from typing import Optional

from app.ai.gemini_client import call_gemini_json
from app.ai.prompts.resume_extraction import RESUME_EXTRACTION_PROMPT

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Text Extraction Helpers
# ---------------------------------------------------------------------------

def extract_text_from_pdf(file_path: str) -> str:
    """Extract raw text from a PDF file using PyMuPDF."""
    import fitz  # PyMuPDF
    try:
        doc = fitz.open(file_path)
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text("text"))
        doc.close()
        full_text = "\n".join(text_parts).strip()
        logger.info(f"Extracted {len(full_text)} chars from PDF: {file_path}")
        return full_text
    except Exception as e:
        logger.error(f"PDF extraction failed for {file_path}: {e}")
        raise


def extract_text_from_docx(file_path: str) -> str:
    """Extract raw text from a DOCX file using python-docx."""
    import docx
    try:
        doc = docx.Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        full_text = "\n".join(paragraphs)
        logger.info(f"Extracted {len(full_text)} chars from DOCX: {file_path}")
        return full_text
    except Exception as e:
        logger.error(f"DOCX extraction failed for {file_path}: {e}")
        raise


def extract_resume_text(file_path: str) -> str:
    """Auto-detect file type and extract text."""
    path = Path(file_path)
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported resume format: {ext}. Only PDF and DOCX are accepted.")


# ---------------------------------------------------------------------------
# Post-processing Helpers
# ---------------------------------------------------------------------------

def compute_total_experience(experience_list: list[dict]) -> float:
    """
    Compute total years of experience from a list of experience entries.
    Each entry should have 'duration_months' or 'start_year'/'end_year'.
    Returns float years rounded to 1 decimal.
    """
    total_months = 0
    for exp in experience_list:
        if "duration_months" in exp and exp["duration_months"]:
            total_months += int(exp["duration_months"])
        elif "start_year" in exp and exp["start_year"]:
            start = int(exp["start_year"])
            end_raw = exp.get("end_year", "present")
            if str(end_raw).lower() in ("present", "current", "now", ""):
                from datetime import datetime
                end = datetime.now().year
            else:
                try:
                    end = int(end_raw)
                except (ValueError, TypeError):
                    from datetime import datetime
                    end = datetime.now().year
            total_months += max(0, (end - start) * 12)
    return round(total_months / 12, 1)


def normalize_skills(skills: list[str]) -> list[str]:
    """Normalize skill names — lowercase, strip whitespace, deduplicate."""
    SKILL_ALIASES = {
        "js": "javascript",
        "ts": "typescript",
        "py": "python",
        "ml": "machine learning",
        "dl": "deep learning",
        "nlp": "natural language processing",
        "cv": "computer vision",
        "react.js": "react",
        "reactjs": "react",
        "node.js": "node.js",
        "nodejs": "node.js",
        "postgres": "postgresql",
        "mongo": "mongodb",
        "k8s": "kubernetes",
    }
    normalized = []
    seen = set()
    for skill in skills:
        s = skill.strip().lower()
        s = SKILL_ALIASES.get(s, s)
        if s and s not in seen:
            seen.add(s)
            # Re-capitalize properly
            normalized.append(s.title() if len(s) > 3 else s.upper())
    return normalized


def extract_email_from_text(text: str) -> Optional[str]:
    """Fallback email extractor using regex."""
    pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    matches = re.findall(pattern, text)
    return matches[0] if matches else None


def extract_phone_from_text(text: str) -> Optional[str]:
    """Fallback phone extractor using regex."""
    pattern = r'(?:\+91[\s-]?)?(?:0)?[6-9]\d{9}'
    matches = re.findall(pattern, text.replace(" ", "").replace("-", ""))
    return matches[0] if matches else None


# ---------------------------------------------------------------------------
# Main Engine
# ---------------------------------------------------------------------------

async def analyze_resume(file_path: str) -> dict:
    """
    Full resume analysis pipeline:
    1. Extract raw text from PDF/DOCX
    2. Send to Gemini for structured extraction
    3. Post-process and normalize results
    4. Return final structured dict

    Returns:
    {
        "raw_text": str,
        "candidate": {
            "first_name": str,
            "last_name": str,
            "email": str,
            "phone": str,
            "linkedin_url": str | None,
            "location": str | None,
        },
        "education": [
            {"degree": str, "institution": str, "year": int | None, "score": str | None}
        ],
        "experience": [
            {
                "company": str,
                "title": str,
                "start_year": int | None,
                "end_year": int | str,
                "duration_months": int | None,
                "key_responsibilities": [str]
            }
        ],
        "skills": [str],  # normalized
        "certifications": [str],
        "total_experience_years": float,
        "summary": str,  # AI-generated 2-line summary
        "extraction_confidence": float  # 0-1
    }
    """
    logger.info(f"Starting resume analysis for: {file_path}")

    # Step 1: Extract text
    raw_text = extract_resume_text(file_path)

    if len(raw_text.strip()) < 100:
        logger.warning(f"Very short resume text ({len(raw_text)} chars) — may be image-based PDF")
        return _empty_result(raw_text, "Resume appears to be image-based or unreadable.")

    # Step 2: Truncate to 6000 chars to stay within Gemini free-tier limits
    truncated_text = raw_text[:6000]

    # Step 3: Call Gemini for structured extraction
    prompt = RESUME_EXTRACTION_PROMPT.format(resume_text=truncated_text)
    try:
        extracted: dict = await call_gemini_json(prompt)
    except Exception as e:
        logger.error(f"Gemini extraction failed: {e}")
        # Fallback: return partial data from regex
        return _fallback_result(raw_text)

    # Step 4: Post-process
    skills_raw = extracted.get("skills", [])
    if isinstance(skills_raw, str):
        skills_raw = [s.strip() for s in skills_raw.split(",")]
    skills_normalized = normalize_skills(skills_raw)

    experience_list = extracted.get("experience", [])
    total_exp = compute_total_experience(experience_list)

    # Ensure email/phone have fallbacks
    candidate = extracted.get("candidate", {})
    if not candidate.get("email"):
        candidate["email"] = extract_email_from_text(raw_text)
    if not candidate.get("phone"):
        candidate["phone"] = extract_phone_from_text(raw_text)

    result = {
        "raw_text": raw_text,
        "candidate": candidate,
        "education": extracted.get("education", []),
        "experience": experience_list,
        "skills": skills_normalized,
        "certifications": extracted.get("certifications", []),
        "total_experience_years": total_exp,
        "summary": extracted.get("summary", ""),
        "extraction_confidence": _compute_confidence(extracted),
    }

    logger.info(
        f"Resume analysis complete. "
        f"Skills: {len(skills_normalized)}, "
        f"Experience: {total_exp}y, "
        f"Confidence: {result['extraction_confidence']}"
    )
    return result


def _compute_confidence(extracted: dict) -> float:
    """Estimate extraction quality based on completeness of key fields."""
    score = 0.0
    candidate = extracted.get("candidate", {})
    if candidate.get("first_name"):
        score += 0.15
    if candidate.get("email"):
        score += 0.15
    if extracted.get("skills"):
        score += 0.20
    if extracted.get("experience"):
        score += 0.25
    if extracted.get("education"):
        score += 0.15
    if extracted.get("summary"):
        score += 0.10
    return round(min(score, 1.0), 2)


def _empty_result(raw_text: str, reason: str) -> dict:
    return {
        "raw_text": raw_text,
        "candidate": {},
        "education": [],
        "experience": [],
        "skills": [],
        "certifications": [],
        "total_experience_years": 0.0,
        "summary": reason,
        "extraction_confidence": 0.0,
    }


def _fallback_result(raw_text: str) -> dict:
    """Basic fallback using regex when Gemini fails."""
    return {
        "raw_text": raw_text,
        "candidate": {
            "email": extract_email_from_text(raw_text),
            "phone": extract_phone_from_text(raw_text),
        },
        "education": [],
        "experience": [],
        "skills": [],
        "certifications": [],
        "total_experience_years": 0.0,
        "summary": "Auto-extraction failed. Manual review required.",
        "extraction_confidence": 0.1,
    }


# ---------------------------------------------------------------------------
# Engine wrapper class and instance export
# ---------------------------------------------------------------------------

import os
import uuid

class ResumeIntelligenceEngine:
    async def parse_resume(self, file_bytes: bytes, filename: str) -> dict:
        upload_dir = Path("./uploads")
        upload_dir.mkdir(parents=True, exist_ok=True)
        
        temp_path = upload_dir / f"temp_{uuid.uuid4()}_{filename}"
        try:
            with open(temp_path, "wb") as f:
                f.write(file_bytes)
            
            result = await analyze_resume(str(temp_path))
            return result
        finally:
            if temp_path.exists():
                try:
                    os.remove(temp_path)
                except Exception as e:
                    logger.error(f"Error removing temp resume file: {e}")


resume_engine = ResumeIntelligenceEngine()

